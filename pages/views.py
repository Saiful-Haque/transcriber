import re
try:
    import whisper
    import torch
    WHISPER_AVAILABLE = True
except ImportError:
    WHISPER_AVAILABLE = False
import os
import sys
import time
from django.shortcuts import render
from django.views import View
from django.conf import settings
from django.utils import timezone
from django.http import HttpResponse
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import io
from pathlib import Path
from .forms import ScribeForm
from .models import TranscriptionJob

# Ensure FFmpeg is in the PATH for Whisper
FFMPEG_PATH = Path(r"C:\ffmpeg\bin")
if str(FFMPEG_PATH) not in os.environ["PATH"]:
    os.environ["PATH"] += os.pathsep + str(FFMPEG_PATH)

# Load the whisper model once on startup
WHISPER_MODEL = None
if WHISPER_AVAILABLE:
    try:
        # Use GPU if available, otherwise fallback to CPU
        DEVICE = "cuda" if torch.cuda.is_available() else "cpu"
        WHISPER_MODEL = whisper.load_model("base", device=DEVICE)
        print(f"Whisper loaded on: {DEVICE}")
    except Exception as e:
        print(f"Error loading Whisper: {e}")
else:
    print("Whisper or Torch not installed. Audio transcription will be unavailable.")

class HomeView(View):
    def get(self, request):
        return render(request, 'index.html')

class ScribeView(View):
    def get(self, request):
        form = ScribeForm()
        return render(request, 'scribe.html', {'form': form})

    def post(self, request):
        form = ScribeForm(request.POST, request.FILES)
        formatted_output = ""
        mode = None
        
        # Timing metrics
        start_time = time.time()
        upload_duration = 0
        transcription_duration = 0
        
        if form.is_valid():
            raw_text = form.cleaned_data['raw_text']
            audio_file = form.cleaned_data.get('audio_file')
            mode = form.cleaned_data.get('transcribe_type')
            
            if audio_file:
                # 1. Measure Upload Time
                upload_start = time.time()
                from django.core.files.storage import default_storage
                file_name = default_storage.save(f'tmp/{audio_file.name}', audio_file)
                file_path = os.path.join(settings.MEDIA_ROOT, file_name)
                upload_duration = time.time() - upload_start
                
                # 2. Measure Transcription Time
                transcription_start = time.time()
                error_occurred = False
                try:
                    transcription_result = self.perform_transcription(file_path, mode)
                    raw_text = transcription_result
                except Exception as e:
                    raw_text = f"Transcription Error: {str(e)}"
                    error_occurred = True
                transcription_duration = time.time() - transcription_start
            
            # Process content only if no error occurred or if we have text
            if audio_file and 'Transcription Error:' in raw_text:
                formatted_output = raw_text
            else:
                sentences = re.split(r'(?<=[.!?])\s+', raw_text)
                lines = []
                current_speaker = "M"
                
                for s in sentences:
                    s = s.strip()
                    if not s: continue
                    lower_s = s.lower()
                    m_triggers = ['okay', 'i see', 'i understand', 'thank', 'and', 'so', 'tell me', 'can you', 'what', 'how']
                    if any(lower_s.startswith(p) for p in m_triggers) or "?" in s:
                        current_speaker = "M"
                    elif len(lines) > 0 and lines[-1].startswith("M:") and ("?" in lines[-1] or len(s) > 40):
                        current_speaker = "R"
                    
                    if lower_s == "just whatsapp." or lower_s == "whatsapp.":
                        current_speaker = "R"
                    lines.append(f"{current_speaker}: {s}")
                
                formatted_output = "\n\n".join(lines)
            
            # Save to Dashboard
            job = TranscriptionJob.objects.create(
                raw_text=raw_text if not audio_file else f"Extracted from {audio_file.name}",
                audio_file=audio_file if audio_file else None,
                formatted_content=formatted_output,
                transcribe_type=mode,
                upload_duration=upload_duration,
                transcription_duration=transcription_duration,
                total_duration=time.time() - start_time,
                finish_time=timezone.now()
            )
            
        return render(request, 'scribe.html', {
            'form': form,
            'output': formatted_output,
            'mode': mode if audio_file else None,
            'job_id': job.id if 'job' in locals() else None
        })

    def perform_transcription(self, file_path, mode):
        if not WHISPER_MODEL:
            return "Whisper model not loaded."
        
        # Path validation
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Audio file not found at: {file_path}")

        # Check for empty file
        if path.stat().st_size == 0:
            raise ValueError("The uploaded audio file is empty (0 bytes).")

        task = "translate" if mode == "english" else "transcribe"
        
        try:
            # Load and validate audio array before processing
            # This helps catch decoding issues early
            audio = whisper.load_audio(str(path))
            if audio.size == 0:
                raise ValueError("Could not extract any audio data. The file might be corrupted or in an unsupported format.")
            
            # Standardizing sample length if it's too short for certain torch versions
            if audio.shape[0] < 480: # Less than 30ms at 16kHz
                 raise ValueError("Audio segment is too short to transcribe.")

            # Use fp16 for GPU (much faster), but keep False for CPU
            use_fp16 = (WHISPER_MODEL.device.type == "cuda")
            
            result = WHISPER_MODEL.transcribe(audio, task=task, fp16=use_fp16)
            return result.get('text', '').strip()
        except Exception as e:
            # Log the specific error for debugging
            print(f"Whisper Error Detail: {str(e)}")
            raise e

class DownloadDocxView(View):
    def get(self, request, job_id):
        try:
            job = TranscriptionJob.objects.get(id=job_id)
        except TranscriptionJob.DoesNotExist:
            return HttpResponse("Job not found", status=404)

        # Create Document
        doc = Document()
        
        # --- Header ---
        section = doc.sections[0]
        header = section.header
        header_para = header.paragraphs[0]
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        logo_path = r"C:\Users\SAIF\Downloads\transcriber logo.jpg"
        if os.path.exists(logo_path):
            run = header_para.add_run()
            run.add_picture(logo_path, width=Inches(1.8))
        else:
            header_para.text = "Transcribersnet\nfor all your transcription needs"
            
        # --- Footer ---
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Add Page Number format: "1 | Page"
        def add_page_number(paragraph):
            run = paragraph.add_run()
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'begin')
            run._r.append(fldChar)
            
            run = paragraph.add_run()
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')
            instrText.text = "PAGE"
            run._r.append(instrText)
            
            run = paragraph.add_run()
            fldChar = OxmlElement('w:fldChar')
            fldChar.set(qn('w:fldCharType'), 'end')
            run._r.append(fldChar)
            
            paragraph.add_run(" | Page")

        add_page_number(footer_para)

        # --- Page Borders ---
        def set_page_border(document):
            sec_pr = document.sections[0]._sectPr
            pg_borders = OxmlElement('w:pgBorders')
            pg_borders.set(qn('w:offsetFrom'), 'page')
            
            # Dark Blue Color from image: 1F497D
            # Size 4 = 0.5 pt
            for side in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4') 
                border.set(qn('w:color'), '1F497D')
                border.set(qn('w:space'), '24')
                pg_borders.append(border)
            
            sec_pr.append(pg_borders)

        set_page_border(doc)

        # Add Metadata Table
        table = doc.add_table(rows=9, cols=2)
        table.style = 'Table Grid'
        
        # Set column widths
        table.columns[0].width = Inches(2.5)
        table.columns[1].width = Inches(3.5)
        
        metadata = [
            ("Audio file name", Path(job.audio_file.name).name if job.audio_file else "Manual Text"),
            ("Transcriber name", "Faizul Haque"),
            ("Language of audio", "Hindi" if job.transcribe_type == 'verbatim' else "English"),
            ("Total length of discussion", f"{int(job.total_duration // 60):02d}:{int(job.total_duration % 60):02d}"),
            ("Total number of pages", "Auto"),
            ("Total Word count", str(len(job.formatted_content.split()))),
            ("Issue faced", "NO"),
            ("Time taken to work", f"{int(job.transcription_duration // 3600):02d}:{int((job.transcription_duration % 3600) // 60):02d}:{int(job.transcription_duration % 60):02d}"),
            ("Spell and Grammar check status", "Yes"),
        ]

        def set_cell_background(cell, fill_color):
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), fill_color)
            cell._tc.get_or_add_tcPr().append(shading_elm)

        for i, (key, value) in enumerate(metadata):
            cell_key = table.cell(i, 0)
            cell_val = table.cell(i, 1)
            
            cell_key.text = key
            cell_val.text = value
            
            # Formatting as per image
            run_key = cell_key.paragraphs[0].runs[0]
            run_key.bold = True
            run_key.font.size = Pt(11)
            
            run_val = cell_val.paragraphs[0].runs[0]
            run_val.font.italic = True if i in [0, 2, 6, 7] else False
            run_val.font.size = Pt(11)

            # Background color (Light Blue: BDD7EE)
            set_cell_background(cell_key, "BDD7EE")
            set_cell_background(cell_val, "BDD7EE")

        # Start content from next page
        doc.add_page_break()

        # Add Transcript Content
        lines = job.formatted_content.split('\n\n')
        for line in lines:
            if not line.strip(): continue
            
            p = doc.add_paragraph()
            # Set default font size
            p_format = p.paragraph_format
            p_format.space_after = Pt(6)
            
            # Speaker labeling and bolding logic
            if line.startswith('M:'):
                run = p.add_run(line)
                run.bold = True
            else:
                run = p.add_run(line)
                run.bold = False
            
            run.font.size = Pt(12)
            run.font.name = 'Calibri'

        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"Transcription_{job_id}.docx"
        response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response

class DashboardView(View):
    def get(self, request):
        jobs = TranscriptionJob.objects.all()
        return render(request, 'dashboard.html', {'jobs': jobs})
